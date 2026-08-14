from __future__ import annotations

import csv
import io
from pathlib import Path

from atlas.ingestion.models import ParsedAttachment

MAX_ATTACHMENT_CHARS = 80_000
IMAGE_TYPES = {
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/gif",
    "image/webp",
    "image/tiff",
}


def extract_attachment(filename: str, content_type: str, payload: bytes) -> ParsedAttachment:
    name = filename or "unnamed"
    ctype = (content_type or "application/octet-stream").split(";")[0].strip().lower()
    suffix = Path(name).suffix.lower()

    if ctype in IMAGE_TYPES or suffix in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".tif", ".tiff"}:
        return ParsedAttachment(
            filename=name,
            content_type=ctype,
            skipped_reason="image attachment — OCR is not enabled yet",
        )

    try:
        text = _extract_text(name, ctype, suffix, payload)
    except Exception as exc:
        return ParsedAttachment(
            filename=name,
            content_type=ctype,
            skipped_reason=f"could not read attachment: {exc}",
        )

    text = (text or "").strip()
    if not text:
        return ParsedAttachment(
            filename=name,
            content_type=ctype,
            skipped_reason="no extractable text (scanned PDF/image needs OCR)",
        )
    if len(text) > MAX_ATTACHMENT_CHARS:
        text = text[:MAX_ATTACHMENT_CHARS] + "\n[truncated]"
    return ParsedAttachment(filename=name, content_type=ctype, text=text)


def _extract_text(name: str, ctype: str, suffix: str, payload: bytes) -> str:
    if ctype == "application/pdf" or suffix == ".pdf":
        return _pdf_text(payload)
    if (
        ctype
        in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }
        or suffix == ".docx"
    ):
        return _docx_text(payload)
    if (
        ctype
        in {
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        }
        or suffix in {".xlsx", ".xlsm"}
    ):
        return _xlsx_text(payload)
    if ctype in {"text/csv", "application/csv"} or suffix == ".csv":
        return payload.decode("utf-8", errors="replace")
    if ctype in {"text/plain", "text/markdown"} or suffix in {".txt", ".md", ".log"}:
        return payload.decode("utf-8", errors="replace")
    if ctype == "text/html" or suffix in {".html", ".htm"}:
        from atlas.ingestion.parse import html_to_text

        return html_to_text(payload.decode("utf-8", errors="replace"))
    if suffix in {".csv"}:
        return payload.decode("utf-8", errors="replace")
    return ""


def _pdf_text(payload: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(payload))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        extracted = page.extract_text() or ""
        if extracted.strip():
            pages.append(f"[PDF page {index}]\n{extracted.strip()}")
    return "\n\n".join(pages)


def _docx_text(payload: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(payload))
    parts = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _xlsx_text(payload: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(payload), read_only=True, data_only=True)
    blocks: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if index > 400:
                rows.append("[truncated extra spreadsheet rows]")
                break
            values = ["" if cell is None else str(cell) for cell in row]
            if any(value.strip() for value in values):
                buffer = io.StringIO()
                csv.writer(buffer).writerow(values)
                rows.append(buffer.getvalue().rstrip("\r\n"))
        if rows:
            blocks.append(f"[Sheet: {sheet.title}]\n" + "\n".join(rows))
    return "\n\n".join(blocks)
